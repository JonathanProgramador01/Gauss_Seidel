from docx import Document
from docx.shared import Pt
import roman


class Word():
    def __init__(self, lista, heading):
        self.lista = lista
        self.document = Document()
        self.document.add_heading(heading, level=0)
        font = self.document.paragraphs[0].runs[0].font
        font.size = Pt(50)
        self.p = self.document.add_paragraph('')
        self.p.add_run("Ecuaciones dadas:").bold = True
        font = self.p.runs[0].font
        font.size = Pt(20)
        self.imprimir()


    def return_lista (self):
        return self.lista
    def cerrar_documento(self,text):
        self.document.save(f"{text}_Jonathan_Rosas_Cedillo.docx")

    def imprimir(self):
        for filas in self.lista:
            self.p = self.document.add_paragraph(f"{filas[0]}x   {filas[1]}y   {filas[2]}z   =   {filas[3]}")
            font = self.p.runs[0].font
            font.size = Pt(12)


    def despeje(self,):
        self.p = self.document.add_paragraph('')
        self.p.add_run("\nDespeje: ").bold = True
        font = self.p.runs[0].font
        font.size = Pt(20)

        self.p = self.document.add_paragraph(f"x = ({self.lista[0][3]}   {self.lista[0][1] * -1}   {self.lista[0][2] * -1}) / {self.lista[0][0]}")
        font = self.p.runs[0].font
        font.size = Pt(12)
        self.p = self.document.add_paragraph(f"y = ({self.lista[1][3]}   {self.lista[1][0] * -1}   {self.lista[1][2] * -1}) / {self.lista[1][1]}")
        font = self.p.runs[0].font
        font.size = Pt(12)
        self.p = self.document.add_paragraph(f"z = ({self.lista[2][3]}   {self.lista[2][0] * -1}   {self.lista[2][1] * -1}) / {self.lista[2][2]}")
        font = self.p.runs[0].font
        font.size = Pt(12)
        self.p = self.document.add_paragraph('')
        self.p.add_run("\nIteraciones:").bold = True
        font = self.p.runs[0].font
        font.size = Pt(20)

    def checar_si_esta_acomodada_tu_ecuacion(self,flag):
        if flag == 0:
            self.p = self.document.add_paragraph("")
            self.p.add_run("\nPerfecto ya cuentas con tus Ecuaciones ordenadas\n").bold = True
            font = self.p.runs[0].font
            font.size = Pt(18)
        else:
            self.p = self.document.add_paragraph('')
            self.p.add_run("\nAcomodando tus Ecuaciones: ").bold = True
            font = self.p.runs[0].font
            font.size = Pt(20)
            self.imprimir()


    def tabla(self,iteraciones,num):
        table = self.document.add_table(rows=1,cols=4)
        top = table.rows[0].cells
        top[0].text = ""
        top[1].text = "X"
        top[2].text = "Y"
        top[3].text = "Z"

        for i in range(num):
            row = table.add_row().cells
            row[0].text = f"{roman.toRoman(i+1)}"
            row[1].text = "{:.4f}".format(iteraciones["X"][i])
            row[2].text = "{:.4f}".format(iteraciones["Y"][i])
            row[3].text = "{:.4f}".format(iteraciones["Z"][i])



















